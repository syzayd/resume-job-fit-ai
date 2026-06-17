"""Streamlit UI for the AI Resume / Job-Fit Tool.

Run: streamlit run app.py
"""

import html as _html
import io
import os
from pathlib import Path

import streamlit as st

# Streamlit Community Cloud stores secrets in st.secrets, not env vars.
# Inject them into os.environ so analyzer.py's os.environ.get() calls work on Cloud.
for _secret_key in ("GEMINI_API_KEY", "GOOGLE_API_KEY"):
    if _secret_key in st.secrets and not os.environ.get(_secret_key):
        os.environ[_secret_key] = st.secrets[_secret_key]

from analyzer import (
    Analysis, AnalyzerError, CoverLetter, InterviewPrep, LinkedInProfile, SkillsRoadmap,
    analyze, generate_cover_letter, generate_interview_prep, generate_linkedin_profile,
    generate_skills_roadmap,
)
from db import save_application

SAMPLE_DIR = Path(__file__).parent / "sample"

st.set_page_config(page_title="Resume Job-Fit AI", page_icon="🎯", layout="wide")


# --- Helpers -----------------------------------------------------------------

def load_sample(name: str) -> str:
    try:
        return (SAMPLE_DIR / name).read_text(encoding="utf-8")
    except OSError:
        return ""


def extract_pdf_text(uploaded_file) -> str:
    import pdfplumber
    with pdfplumber.open(io.BytesIO(uploaded_file.read())) as pdf:
        return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()


def score_color(score: int) -> str:
    if score >= 75:
        return "#16a34a"
    if score >= 50:
        return "#d97706"
    return "#dc2626"


def chips(items: list[str], bg: str, fg: str) -> None:
    if not items:
        st.caption("None found.")
        return
    html = " ".join(
        f"<span style='background:{bg};color:{fg};padding:4px 10px;border-radius:14px;"
        f"font-size:0.85rem;margin:2px;display:inline-block;'>{_html.escape(item)}</span>"
        for item in items
    )
    st.markdown(html, unsafe_allow_html=True)


def analysis_as_text(
    result: Analysis,
    cover: CoverLetter | None = None,
    prep: InterviewPrep | None = None,
    roadmap: SkillsRoadmap | None = None,
    linkedin: LinkedInProfile | None = None,
) -> str:
    lines = [
        "RESUME JOB-FIT ANALYSIS",
        "=" * 40,
        f"Score: {result.score}/100",
        f"Verdict: {result.verdict}",
    ]
    if result.salary_range:
        lines += [f"Estimated Salary Range: {result.salary_range}"]
    lines += [
        "",
        "MATCHED KEYWORDS",
        ", ".join(result.matched_keywords) or "None",
        "",
        "MISSING KEYWORDS",
        ", ".join(result.missing_keywords) or "None",
        "",
        "BULLET REWRITES",
    ]
    for rw in result.bullet_rewrites:
        lines += [f"  Before: {rw.original}", f"  After:  {rw.improved}", ""]
    lines += ["", "HOW TO IMPROVE", result.summary, "", "ATS TIPS"]
    lines += [f"  - {tip}" for tip in result.ats_tips]
    if cover:
        lines += ["", "=" * 40, "COVER LETTER", "=" * 40, "",
                  cover.opening, "", cover.body, "", cover.closing]
    if prep:
        lines += ["", "=" * 40, "INTERVIEW PREP", "=" * 40, ""]
        lines += [f"Key advice: {prep.opening_tip}", ""]
        for i, q in enumerate(prep.questions, 1):
            lines += [f"Q{i}: {q.question}", f"    Why asked: {q.why_asked}",
                      f"    Tip: {q.tip}", ""]
    if roadmap:
        lines += ["", "=" * 40, "SKILLS ROADMAP", "=" * 40, "",
                  f"Timeline: {roadmap.timeline}", "",
                  "Quick wins this week:"]
        lines += [f"  - {w}" for w in roadmap.quick_wins]
        lines += ["", "Skill gaps (priority order):"]
        for gap in roadmap.gaps:
            lines += [f"\n  [{gap.importance}] {gap.skill}", f"  {gap.how_to_learn}"]
            for r in gap.resources:
                lines += [f"    - {r.name} ({r.provider}, {r.type})"]
    if linkedin:
        lines += ["", "=" * 40, "LINKEDIN PROFILE", "=" * 40, "",
                  f"Headline: {linkedin.headline}", "",
                  "ABOUT SECTION", linkedin.about, "",
                  "SKILLS TO ADD", ", ".join(linkedin.skills_to_add), "",
                  "PROFILE TIPS"]
        lines += [f"  - {tip}" for tip in linkedin.profile_tips]
    return "\n".join(lines)


def export_docx(
    result: Analysis,
    cover: CoverLetter | None = None,
    prep: InterviewPrep | None = None,
    roadmap: SkillsRoadmap | None = None,
    linkedin: LinkedInProfile | None = None,
) -> bytes:
    """Build a formatted Word document and return it as bytes."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Resume Job-Fit Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Score
    score_para = doc.add_paragraph()
    score_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = score_para.add_run(f"Fit Score: {result.score}/100")
    run.bold = True
    run.font.size = Pt(18)
    doc.add_paragraph(result.verdict)
    if result.salary_range:
        p = doc.add_paragraph()
        p.add_run("Estimated Salary Range: ").bold = True
        p.add_run(result.salary_range)

    # Keywords
    doc.add_heading("Matched Keywords", 2)
    doc.add_paragraph(", ".join(result.matched_keywords) or "None")

    doc.add_heading("Missing Keywords", 2)
    doc.add_paragraph(", ".join(result.missing_keywords) or "None")

    # Bullet rewrites
    doc.add_heading("Tailored Bullet Rewrites", 2)
    for rw in result.bullet_rewrites:
        p = doc.add_paragraph()
        p.add_run("Before: ").bold = True
        p.add_run(rw.original)
        p2 = doc.add_paragraph()
        p2.add_run("After:  ").bold = True
        p2.add_run(rw.improved)
        doc.add_paragraph()

    # Summary + ATS
    doc.add_heading("How to Improve This Application", 2)
    doc.add_paragraph(result.summary)
    doc.add_heading("ATS Tips", 3)
    for tip in result.ats_tips:
        doc.add_paragraph(tip, style="List Bullet")

    # Cover letter
    if cover:
        doc.add_page_break()
        doc.add_heading("Cover Letter", 1)
        doc.add_paragraph(cover.opening)
        doc.add_paragraph(cover.body)
        doc.add_paragraph(cover.closing)

    # Interview prep
    if prep:
        doc.add_page_break()
        doc.add_heading("Interview Prep", 1)
        doc.add_paragraph(prep.opening_tip)
        for i, q in enumerate(prep.questions, 1):
            doc.add_heading(f"Q{i}: {q.question}", 3)
            p = doc.add_paragraph()
            p.add_run("Why asked: ").bold = True
            p.add_run(q.why_asked)
            p2 = doc.add_paragraph()
            p2.add_run("Tip: ").bold = True
            p2.add_run(q.tip)

    # Skills roadmap
    if roadmap:
        doc.add_page_break()
        doc.add_heading("Skills Gap Roadmap", 1)
        doc.add_paragraph(f"Estimated timeline: {roadmap.timeline}")
        doc.add_heading("Quick Wins This Week", 2)
        for win in roadmap.quick_wins:
            doc.add_paragraph(win, style="List Bullet")
        doc.add_heading("Skill Gaps — Priority Order", 2)
        for gap in roadmap.gaps:
            doc.add_heading(f"[{gap.importance}] {gap.skill}", 3)
            doc.add_paragraph(gap.how_to_learn)
            for r in gap.resources:
                doc.add_paragraph(f"{r.name} — {r.provider} ({r.type})", style="List Bullet")

    # LinkedIn
    if linkedin:
        doc.add_page_break()
        doc.add_heading("LinkedIn Profile Optimizer", 1)
        doc.add_heading("Headline", 2)
        doc.add_paragraph(linkedin.headline)
        doc.add_heading("About Section", 2)
        doc.add_paragraph(linkedin.about)
        doc.add_heading("Skills to Add", 2)
        doc.add_paragraph(", ".join(linkedin.skills_to_add))
        doc.add_heading("Profile Tips", 2)
        for tip in linkedin.profile_tips:
            doc.add_paragraph(tip, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# --- Result panels -----------------------------------------------------------

def render_analysis(result: Analysis) -> None:
    color = score_color(result.score)
    st.markdown(
        f"<div style='text-align:center;margin:0.5rem 0;'>"
        f"<span style='font-size:3.5rem;font-weight:800;color:{color};'>{result.score}</span>"
        f"<span style='font-size:1.2rem;color:#888;'>/100</span></div>",
        unsafe_allow_html=True,
    )
    st.markdown(
        f"<p style='text-align:center;font-size:1.1rem;'>{_html.escape(result.verdict)}</p>",
        unsafe_allow_html=True,
    )
    if result.salary_range:
        st.markdown(
            f"<div style='text-align:center;margin:0.25rem 0 0.75rem;'>"
            f"<span style='background:#fefce8;color:#854d0e;border:1px solid #fde68a;"
            f"padding:5px 14px;border-radius:20px;font-size:0.9rem;font-weight:600;'>"
            f"💰 Estimated salary: {_html.escape(result.salary_range)}</span></div>",
            unsafe_allow_html=True,
        )
    st.divider()

    left, right = st.columns(2)
    with left:
        st.subheader("Matched keywords")
        chips(result.matched_keywords, "#dcfce7", "#166534")
    with right:
        st.subheader("Missing keywords")
        chips(result.missing_keywords, "#fee2e2", "#991b1b")

    st.divider()
    st.subheader("Tailored bullet rewrites")
    st.caption("Original on the left · AI rewrite on the right · Copy individual rewrites or all at once.")
    for i, rw in enumerate(result.bullet_rewrites):
        left_col, right_col = st.columns(2)
        with left_col:
            st.markdown(
                f"<div style='background:#f3f4f6;border-left:3px solid #9ca3af;"
                f"padding:12px 14px;border-radius:6px;font-size:0.88rem;'>"
                f"<div style='color:#6b7280;font-size:0.72rem;font-weight:700;"
                f"text-transform:uppercase;letter-spacing:0.06em;margin-bottom:8px;'>Original</div>"
                f"{_html.escape(rw.original)}</div>",
                unsafe_allow_html=True,
            )
        with right_col:
            st.markdown(
                f"<div style='background:#f0fdf4;border-left:3px solid #16a34a;"
                f"padding:12px 14px;border-radius:6px;font-size:0.88rem;'>"
                f"<div style='color:#16a34a;font-size:0.72rem;font-weight:700;"
                f"text-transform:uppercase;letter-spacing:0.06em;margin-bottom:8px;'>Rewritten ✓</div>"
                f"{_html.escape(rw.improved)}</div>",
                unsafe_allow_html=True,
            )
            with st.expander(f"Copy rewrite #{i + 1}"):
                st.code(rw.improved, language=None)
        st.write("")
    if result.bullet_rewrites:
        with st.expander("Copy all rewrites"):
            st.code(
                "\n\n".join(rw.improved for rw in result.bullet_rewrites),
                language=None,
            )

    st.divider()
    st.subheader("How to improve this application")
    st.write(result.summary)
    if result.ats_tips:
        st.markdown("**ATS tips:**")
        for tip in result.ats_tips:
            st.markdown(f"- {tip}")


def render_cover_letter(cover: CoverLetter) -> None:
    st.write(cover.opening)
    st.write(cover.body)
    st.write(cover.closing)
    st.divider()
    with st.expander("Copy full letter"):
        st.code(f"{cover.opening}\n\n{cover.body}\n\n{cover.closing}", language=None)


def render_interview_prep(prep: InterviewPrep) -> None:
    st.info(f"**Key advice:** {prep.opening_tip}")
    st.divider()
    for i, q in enumerate(prep.questions, 1):
        with st.container(border=True):
            st.markdown(f"**Q{i}: {q.question}**")
            st.caption(f"Why asked: {q.why_asked}")
            st.markdown(f"Tip: {q.tip}")
            with st.expander("Copy question + tip"):
                st.code(f"Q: {q.question}\n\nTip: {q.tip}", language=None)


def render_skills_roadmap(roadmap: SkillsRoadmap) -> None:
    st.markdown(f"**Estimated timeline to close key gaps:** {roadmap.timeline}")
    if roadmap.quick_wins:
        st.subheader("Quick wins this week")
        for win in roadmap.quick_wins:
            st.markdown(f"- {win}")
    st.divider()
    st.subheader("Skill gaps — priority order")
    importance_color = {"High": "#fee2e2", "Medium": "#fef9c3", "Low": "#f0fdf4"}
    importance_fg = {"High": "#991b1b", "Medium": "#854d0e", "Low": "#166534"}
    for gap in roadmap.gaps:
        bg = importance_color.get(gap.importance, "#f3f4f6")
        fg = importance_fg.get(gap.importance, "#374151")
        with st.container(border=True):
            st.markdown(
                f"<span style='background:{bg};color:{fg};padding:2px 8px;"
                f"border-radius:10px;font-size:0.8rem;font-weight:600;'>{_html.escape(gap.importance)}</span>"
                f" &nbsp; **{_html.escape(gap.skill)}**",
                unsafe_allow_html=True,
            )
            st.write(gap.how_to_learn)
            if gap.resources:
                for r in gap.resources:
                    st.markdown(f"- **{r.name}** — {r.provider} _{r.type}_")


def render_linkedin_profile(profile: LinkedInProfile) -> None:
    st.subheader("Headline")
    st.markdown(f"> {_html.escape(profile.headline)}", unsafe_allow_html=False)
    with st.expander("Copy headline"):
        st.code(profile.headline, language=None)

    st.divider()
    st.subheader("About section")
    st.write(profile.about)
    with st.expander("Copy About section"):
        st.code(profile.about, language=None)

    st.divider()
    st.subheader("Skills to add on LinkedIn")
    chips(profile.skills_to_add, "#dbeafe", "#1e40af")

    st.divider()
    st.subheader("Profile tips")
    for tip in profile.profile_tips:
        st.markdown(f"- {tip}")


# --- Session state init ------------------------------------------------------

def _init() -> None:
    defaults = {
        "resume": "",
        "job": "",
        "pdf_name": None,
        "result": None,
        "cover_letter": None,
        "cover_letter_tone": "Professional",
        "interview_prep": None,
        "skills_roadmap": None,
        "linkedin_profile": None,
        "tracker_saved": False,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

_init()

# --- Layout ------------------------------------------------------------------

st.title("Resume Job-Fit AI")
st.caption(
    "Paste or upload your resume + a job description — get a fit score, keyword gaps, "
    "tailored rewrites, cover letter, interview prep, and a skills roadmap. "
    "Powered by Google Gemini (free tier)."
)

col_a, col_b = st.columns(2)
with col_a:
    st.session_state.job = st.text_area(
        "Job description", value=st.session_state.job,
        height=300, placeholder="Paste the job posting here...",
    )
with col_b:
    uploaded_pdf = st.file_uploader(
        "Upload resume PDF (or paste below)", type="pdf", label_visibility="visible",
    )
    if uploaded_pdf is not None and uploaded_pdf.name != st.session_state.pdf_name:
        try:
            st.session_state.resume = extract_pdf_text(uploaded_pdf)
            st.session_state.pdf_name = uploaded_pdf.name
            st.session_state.result = None
            st.session_state.cover_letter = None
            st.session_state.interview_prep = None
            st.session_state.skills_roadmap = None
            st.session_state.linkedin_profile = None
            st.rerun()
        except Exception:
            st.error("Could not read the PDF. Try a text-based PDF or paste your resume below.")

    st.session_state.resume = st.text_area(
        "Your resume", value=st.session_state.resume,
        height=220, placeholder="Paste your resume text here...",
        label_visibility="collapsed",
    )

btn_analyze, btn_sample, btn_clear, _ = st.columns([1, 1, 1, 3])
with btn_analyze:
    analyze_clicked = st.button("Analyze fit", type="primary", use_container_width=True)
with btn_sample:
    if st.button("Load sample", use_container_width=True):
        st.session_state.resume = load_sample("sample_resume.txt")
        st.session_state.job = load_sample("sample_job.txt")
        st.session_state.pdf_name = None
        st.session_state.result = None
        st.session_state.cover_letter = None
        st.session_state.interview_prep = None
        st.session_state.skills_roadmap = None
        st.session_state.linkedin_profile = None
        st.rerun()
with btn_clear:
    if st.button("Clear", use_container_width=True):
        st.session_state.resume = ""
        st.session_state.job = ""
        st.session_state.pdf_name = None
        st.session_state.result = None
        st.session_state.cover_letter = None
        st.session_state.interview_prep = None
        st.session_state.skills_roadmap = None
        st.session_state.linkedin_profile = None
        st.rerun()

if analyze_clicked:
    with st.spinner("Analyzing with Gemini..."):
        try:
            st.session_state.result = analyze(st.session_state.resume, st.session_state.job)
            st.session_state.cover_letter = None
            st.session_state.interview_prep = None
            st.session_state.skills_roadmap = None
            st.session_state.linkedin_profile = None
            st.session_state.tracker_saved = False
        except AnalyzerError as err:
            st.error(str(err))

# --- Results (persist across reruns) -----------------------------------------

if st.session_state.result:
    result: Analysis = st.session_state.result

    # "Generate all sections" — runs all 4 AI generators in one click
    _todo = [k for k in ("cover_letter", "interview_prep", "skills_roadmap", "linkedin_profile")
             if st.session_state[k] is None]
    if _todo:
        _gen_all_col, _ = st.columns([2, 5])
        with _gen_all_col:
            if st.button("Generate all sections ✨", type="primary", use_container_width=True):
                _generators = [
                    (generate_interview_prep,  "interview_prep",  "Interview Prep"),
                    (generate_skills_roadmap,  "skills_roadmap",  "Skills Roadmap"),
                    (generate_linkedin_profile,"linkedin_profile","LinkedIn Profile"),
                ]
                _progress = st.progress(0, text="Starting…")
                if st.session_state.cover_letter is None:
                    _progress.progress(0, text="Generating Cover Letter…")
                    try:
                        st.session_state.cover_letter = generate_cover_letter(
                            st.session_state.resume,
                            st.session_state.job,
                            st.session_state.cover_letter_tone,
                        )
                    except AnalyzerError as _err:
                        st.warning(f"Cover Letter: {_err}")
                for _i, (_fn, _key, _label) in enumerate(_generators):
                    if st.session_state[_key] is None:
                        _progress.progress((_i + 1) / (len(_generators) + 1), text=f"Generating {_label}…")
                        try:
                            st.session_state[_key] = _fn(
                                st.session_state.resume, st.session_state.job
                            )
                        except AnalyzerError as _err:
                            st.warning(f"{_label}: {_err}")
                _progress.progress(1.0, text="All sections ready!")
                st.rerun()

    tab_analysis, tab_cover, tab_interview, tab_roadmap, tab_linkedin = st.tabs(
        ["Analysis", "Cover Letter", "Interview Prep", "Skills Roadmap", "LinkedIn Profile"]
    )

    with tab_analysis:
        render_analysis(result)

    with tab_cover:
        cover: CoverLetter | None = st.session_state.cover_letter
        _tone = st.radio(
            "Tone",
            ["Professional", "Warm & Enthusiastic", "Bold & Direct"],
            index=["Professional", "Warm & Enthusiastic", "Bold & Direct"].index(
                st.session_state.cover_letter_tone
            ),
            horizontal=True,
            key="_tone_radio",
        )
        if _tone != st.session_state.cover_letter_tone:
            st.session_state.cover_letter_tone = _tone
        _btn_label = "Generate cover letter" if cover is None else "Regenerate with this tone"
        if st.button(_btn_label, type="primary"):
            with st.spinner("Writing your cover letter with Gemini..."):
                try:
                    cover = generate_cover_letter(
                        st.session_state.resume, st.session_state.job, _tone
                    )
                    st.session_state.cover_letter = cover
                    st.rerun()
                except AnalyzerError as err:
                    st.error(str(err))
        if cover:
            render_cover_letter(cover)

    with tab_interview:
        prep: InterviewPrep | None = st.session_state.interview_prep
        if prep is None:
            if st.button("Generate interview prep", type="primary"):
                with st.spinner("Generating interview questions with Gemini..."):
                    try:
                        prep = generate_interview_prep(
                            st.session_state.resume, st.session_state.job
                        )
                        st.session_state.interview_prep = prep
                        st.rerun()
                    except AnalyzerError as err:
                        st.error(str(err))
        if prep:
            render_interview_prep(prep)

    with tab_roadmap:
        roadmap: SkillsRoadmap | None = st.session_state.skills_roadmap
        if roadmap is None:
            if st.button("Generate skills roadmap", type="primary"):
                with st.spinner("Building your skills roadmap with Gemini..."):
                    try:
                        roadmap = generate_skills_roadmap(
                            st.session_state.resume, st.session_state.job
                        )
                        st.session_state.skills_roadmap = roadmap
                        st.rerun()
                    except AnalyzerError as err:
                        st.error(str(err))
        if roadmap:
            render_skills_roadmap(roadmap)

    with tab_linkedin:
        li_profile: LinkedInProfile | None = st.session_state.linkedin_profile
        if li_profile is None:
            st.caption(
                "Get an optimized LinkedIn headline, a ready-to-paste About section, "
                "skills to add, and profile tips — all tailored to this role."
            )
            if st.button("Optimize LinkedIn profile", type="primary"):
                with st.spinner("Optimizing your LinkedIn profile with Gemini..."):
                    try:
                        li_profile = generate_linkedin_profile(
                            st.session_state.resume, st.session_state.job
                        )
                        st.session_state.linkedin_profile = li_profile
                        st.rerun()
                    except AnalyzerError as err:
                        st.error(str(err))
        if li_profile:
            render_linkedin_profile(li_profile)

    st.divider()
    _dl_txt, _dl_docx, _ = st.columns([1, 1, 4])
    with _dl_txt:
        st.download_button(
            label="Download (.txt)",
            data=analysis_as_text(
                result,
                st.session_state.cover_letter,
                st.session_state.interview_prep,
                st.session_state.skills_roadmap,
                st.session_state.linkedin_profile,
            ),
            file_name="resume_analysis.txt",
            mime="text/plain",
            use_container_width=True,
        )
    with _dl_docx:
        st.download_button(
            label="Download (.docx)",
            data=export_docx(
                result,
                st.session_state.cover_letter,
                st.session_state.interview_prep,
                st.session_state.skills_roadmap,
                st.session_state.linkedin_profile,
            ),
            file_name="resume_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    # Save to tracker
    st.divider()
    if st.session_state.tracker_saved:
        st.success("Saved to Job Tracker! View it in the **Job Tracker** page (sidebar).")
    else:
        with st.form("save_to_tracker"):
            _tc1, _tc2, _tc3 = st.columns([2, 2, 1])
            with _tc1:
                _job_title = st.text_input(
                    "Job title",
                    placeholder="e.g. Senior Python Engineer",
                )
            with _tc2:
                _company = st.text_input(
                    "Company (optional)",
                    placeholder="e.g. Acme Corp",
                )
            with _tc3:
                st.markdown("<br>", unsafe_allow_html=True)
                _save_clicked = st.form_submit_button(
                    "Save to tracker 📋", use_container_width=True
                )
            if _save_clicked:
                if not _job_title.strip():
                    st.warning("Enter a job title to save.")
                else:
                    save_application(
                        job_title=_job_title,
                        score=result.score,
                        company=_company,
                    )
                    st.session_state.tracker_saved = True
                    st.rerun()

st.divider()
st.caption(
    "Built in public by Zaid Ali Syed "
    "· [Live demo](https://resume-job-fit-ai.streamlit.app) "
    "· [GitHub](https://github.com/syzayd/resume-job-fit-ai) "
    "· Rewrites stay truthful to your resume — review before using."
)
